import os
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime

class Reporter:
    @staticmethod
    def generate_docx(vulnerabilities, output_file="vulnerability_report.docx"):
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('웹 취약점 자동 스캐닝 결과 보고서', 0)
            doc.add_paragraph(f"진단 일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            if not vulnerabilities:
                doc.add_paragraph("\n✅ 발견된 취약점이 없습니다. 안전합니다.")
            else:
                doc.add_heading(f"요약: 총 {len(vulnerabilities)}개의 취약점 발견", level=1)
                
                for idx, vuln in enumerate(vulnerabilities, 1):
                    # 취약점 제목 추가 (예: 1. SQL Injection (Critical))
                    heading = doc.add_heading(f"{idx}. {vuln.title} ({vuln.severity})", level=2)
                    
                    # 심각도별 색상 적용
                    run = heading.runs[0]
                    if vuln.severity == "Critical":
                        run.font.color.rgb = RGBColor(255, 0, 0) # Red
                    elif vuln.severity == "High":
                        run.font.color.rgb = RGBColor(255, 102, 0) # Orange
                    
                    doc.add_paragraph(f"대상 URL: {vuln.endpoint.url}")
                    doc.add_paragraph(f"HTTP 메서드: {vuln.endpoint.method}")
                    doc.add_paragraph(f"영향받은 파라미터: {list(vuln.endpoint.parameters.keys())}")
                    
                    doc.add_paragraph("사용된 페이로드 (Payload Used):", style='List Bullet')
                    payload_para = doc.add_paragraph()
                    r = payload_para.add_run(vuln.payload_used)
                    r.font.name = 'Courier New'
                    r.font.color.rgb = RGBColor(180, 0, 0)
                    
                    doc.add_paragraph("증거 (Evidence):", style='List Bullet')
                    ev_para = doc.add_paragraph()
                    er = ev_para.add_run(vuln.evidence)
                    er.font.name = 'Courier New'
                    
                    doc.add_paragraph("-" * 50)
            
            doc.save(output_file)
            print(f"\n[+] 워드 보고서가 성공적으로 저장되었습니다: {os.path.abspath(output_file)}")
        except Exception as e:
            print(f"\n[!] 워드 문서 생성 중 오류가 발생했습니다: {e}")
